from __future__ import annotations

import asyncio
import io
import re

from docx import Document as DocxDocument

from app.infrastructure.ingestion.ocr.ocr_provider import OCRProvider
from app.infrastructure.ingestion.parsed_document import ParsedDocument, ParsedSection
from app.infrastructure.ingestion.parsers.document_parser import DocumentParser

# Regex patterns for Vietnamese legal document structure
_CHUONG_RE = re.compile(r"^(Chương\s+[IVXLCDM\d]+)", re.IGNORECASE)
_DIEU_RE = re.compile(r"^(Điều\s+\d+)[\.\:]?\s*(.*)", re.IGNORECASE)
_KHOAN_RE = re.compile(r"^(\d+)\.\s+", re.IGNORECASE)
_DIEM_RE = re.compile(r"^([a-zđ])\)\s+", re.IGNORECASE)


def _detect_level(text: str) -> tuple[int, str, str] | None:
    """Return (level, section_number, section_title) or None."""
    m = _CHUONG_RE.match(text)
    if m:
        title = text[m.end() :].strip(" .-–—")  # noqa: RUF001
        return 1, m.group(1), title

    m = _DIEU_RE.match(text)
    if m:
        return 2, m.group(1), m.group(2).strip()

    m = _KHOAN_RE.match(text)
    if m:
        return 3, m.group(1) + ".", text[m.end() :].strip()

    m = _DIEM_RE.match(text)
    if m:
        return 4, m.group(1) + ")", text[m.end() :].strip()

    return None


def _parse_sync(content: bytes) -> ParsedDocument:
    doc = DocxDocument(io.BytesIO(content))
    all_text: list[str] = []
    sections: list[ParsedSection] = []
    current: ParsedSection | None = None
    title = ""

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        all_text.append(text)

        detection = _detect_level(text)
        if detection is not None:
            level, num, section_title = detection
            if current is not None:
                sections.append(current)
            current = ParsedSection(
                section_number=num,
                section_title=section_title,
                content="",
                level=level,
                page_number=0,
            )
        elif current is not None:
            current.content += text + "\n"
        elif not title:
            title = text

    if current is not None:
        sections.append(current)

    raw_text = "\n".join(all_text)
    if not title and all_text:
        title = all_text[0]

    return ParsedDocument(
        title=title,
        raw_text=raw_text,
        sections=sections,
        page_count=len(doc.sections),
    )


class DocxParser(DocumentParser):
    def __init__(self, ocr: OCRProvider | None = None) -> None:
        self._ocr = ocr

    @property
    def supported_content_types(self) -> list[str]:
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]

    async def parse(self, content: bytes, filename: str) -> ParsedDocument:
        return await asyncio.to_thread(_parse_sync, content)
